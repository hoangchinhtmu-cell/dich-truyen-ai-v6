
import io, json, re, time, concurrent.futures
from pathlib import Path
import streamlit as st
from docx import Document
from openai import OpenAI

st.set_page_config(page_title="Dịch Truyện AI V6 Cân Bằng", page_icon="📖", layout="wide")

DEFAULT_STYLE = """Văn phong truyện tự nhiên, mượt, dễ đọc như bản dịch tiểu thuyết Việt được biên tập kỹ.
Giữ sắc thái cảm xúc và bối cảnh. Đối thoại tự nhiên, không máy móc.
Không tự ý thêm, bớt hoặc giải thích nội dung.
Giữ nhất quán tên riêng, quan hệ, giới tính, chức danh, thuật ngữ và xưng hô.
Nếu là cổ trang, ưu tiên văn phong cổ trang; không dùng đại từ hiện đại nếu STORY BIBLE không cho phép."""

MODEL_PRO = "deepseek-v4-pro"
MODEL_FLASH = "deepseek-v4-flash"

def read_docx(data):
    doc = Document(io.BytesIO(data))
    return [p.text.strip() for p in doc.paragraphs if p.text.strip()]

def read_txt(data):
    for enc in ("utf-8-sig", "utf-8", "gb18030", "gbk"):
        try:
            return [x.strip() for x in data.decode(enc).splitlines() if x.strip()]
        except UnicodeDecodeError:
            pass
    raise ValueError("Không đọc được TXT. Hãy lưu TXT bằng UTF-8 hoặc GB18030.")

def split_chapters(paragraphs):
    pat = re.compile(
        r"^\s*(第\s*[0-9一二三四五六七八九十百千万]+\s*[章回节卷篇部]|"
        r"Chương\s+\d+(?:\s*[:：.\-].*)?)\s*$", re.I)
    starts = [i for i,p in enumerate(paragraphs) if pat.match(p)]
    if not starts:
        return [{"number":1,"title":"Chương 1","paragraphs":paragraphs}]
    out=[]
    for n,start in enumerate(starts):
        end=starts[n+1] if n+1<len(starts) else len(paragraphs)
        out.append({"number":n+1,"title":f"Chương {n+1}","source_title":paragraphs[start],
                    "paragraphs":paragraphs[start+1:end]})
    return out

def split_smart(items, limit=9000):
    """Chia ưu tiên ở cuối câu/đoạn, không cắt giữa câu nếu có thể."""
    out=[]; cur=[]; size=0
    sentence_re=re.compile(r"(?<=[。！？!?…])\s*")
    for p in items:
        p=p.strip()
        if not p: continue
        # Đoạn bình thường không quá giới hạn -> giữ nguyên.
        if len(p) <= limit:
            if cur and size+len(p)+1 > limit:
                out.append(cur); cur=[]; size=0
            cur.append(p); size += len(p)+1
            continue
        # Đoạn quá dài -> tách theo câu.
        sentences=[s.strip() for s in sentence_re.split(p) if s.strip()]
        for s in sentences:
            if cur and size+len(s)+1 > limit:
                out.append(cur); cur=[]; size=0
            # Một câu cực dài: cắt mềm theo dấu phẩy/chấm phẩy.
            if len(s) > limit:
                pieces=re.split(r"(?<=[，,；;：:])\s*",s)
                for piece in pieces:
                    if cur and size+len(piece)+1 > limit:
                        out.append(cur); cur=[]; size=0
                    cur.append(piece); size += len(piece)+1
            else:
                cur.append(s); size += len(s)+1
    if cur: out.append(cur)
    return out

def call_json(client, model, system, user, *, thinking=False, timeout=180, retries=3):
    last=None
    for attempt in range(retries):
        try:
            kwargs=dict(
                model=model,
                messages=[{"role":"system","content":system},{"role":"user","content":user}],
                stream=False,
                response_format={"type":"json_object"},
                timeout=timeout,
            )
            kwargs["extra_body"]={"thinking":{"type":"enabled" if thinking else "disabled"}}
            r=client.chat.completions.create(**kwargs)
            text=(r.choices[0].message.content or "").strip()
            if not text:
                raise ValueError("DeepSeek trả về nội dung rỗng.")
            try:
                return json.loads(text)
            except json.JSONDecodeError:
                m=re.search(r"\{.*\}",text,re.S)
                if m: return json.loads(m.group(0))
                raise ValueError("DeepSeek trả về JSON không hợp lệ.")
        except Exception as e:
            last=e
            time.sleep(min(2**attempt,6))
    raise RuntimeError(f"API thất bại sau {retries} lần: {last}")

def build_bible(client, chapters, style, status_cb=None):
    # Lấy mẫu trải đều để không chỉ hiểu vài chương đầu.
    samples=[]
    if chapters:
        indexes=sorted(set([0,len(chapters)//4,len(chapters)//2,(3*len(chapters))//4,len(chapters)-1]))
        for i in indexes:
            ch=chapters[i]
            samples.append(f"### {ch['title']}\n"+"\n".join(ch["paragraphs"][:25]))
    text="\n\n".join(samples)[:70000]
    system="""Bạn là biên tập viên tiểu thuyết Trung -> Việt.
Hãy tạo STORY BIBLE dùng cho cả bộ truyện. Không dịch văn bản.
Chỉ ghi thông tin có căn cứ; chưa rõ thì ghi "chưa xác định".
Đặc biệt phải khóa tên, giới tính, quan hệ và XƯNG HÔ theo từng cặp nhân vật.
Nếu truyện cổ trang, không tự ý hiện đại hóa đại từ."""
    user=f"""Phong cách:
{style}

Trả JSON đúng cấu trúc:
{{
"characters":[
 {{"name_cn":"","name_vi":"","aliases":[],"gender":"","role":"","relations":[],
   "pronouns":[{{"with":"","self":"","other":"","when":""}}],"notes":""}}
],
"glossary":[{{"source":"","translation":"","type":"","notes":""}}],
"world_rules":[],
"style_rules":[],
"uncertain_items":[]
}}

MẪU TRUYỆN:
{text}"""
    return call_json(client,MODEL_PRO,system,user,thinking=False,timeout=240)

def merge_memory(mem, upd):
    mem=mem or {}
    for k in ("characters","glossary","world_rules","style_rules","uncertain_items"):
        mem.setdefault(k,[])
    for x in upd.get("characters",[]) or []:
        name=(x.get("name_cn") or x.get("name_vi") or "").strip()
        if not name: continue
        old=next((a for a in mem["characters"] if (a.get("name_cn") or a.get("name_vi"))==name),None)
        if old:
            for k,v in x.items():
                if v not in ("",None,[],{}): old[k]=v
        else: mem["characters"].append(x)
    for x in upd.get("glossary",[]) or []:
        src=str(x.get("source","")).strip()
        if not src: continue
        old=next((a for a in mem["glossary"] if a.get("source")==src),None)
        if old: old.update({k:v for k,v in x.items() if v not in ("",None)})
        else: mem["glossary"].append(x)
    for k in ("world_rules","style_rules","uncertain_items"):
        for x in upd.get(k,[]) or []:
            if x and x not in mem[k]: mem[k].append(x)
    return mem

def relevant_memory(mem, source):
    """Giảm prompt bằng cách lấy nhân vật/thuật ngữ có dấu hiệu xuất hiện trong chunk."""
    text=source
    chars=[]
    for c in mem.get("characters",[]):
        names=[c.get("name_cn",""),c.get("name_vi","")] + (c.get("aliases") or [])
        if any(n and n in text for n in names):
            chars.append(c)
    # Nếu chunk ít tên hoặc văn bản có đại từ nhiều, giữ thêm một số nhân vật quan trọng đầu tiên.
    if not chars:
        chars=mem.get("characters",[])[:12]
    terms=[g for g in mem.get("glossary",[]) if g.get("source") and g["source"] in text]
    return {
        "characters":chars,
        "glossary":terms,
        "world_rules":mem.get("world_rules",[])[-20:],
        "style_rules":mem.get("style_rules",[])[-20:],
    }

def translate_one(client, title, part, mem, style, use_flash=True):
    source="\n".join(f"[P{i+1}] {p}" for i,p in enumerate(part))
    local=relevant_memory(mem,source)
    system="""Bạn là dịch giả tiểu thuyết Trung -> Việt chuyên nghiệp.
Dịch tự nhiên, mượt, đúng bối cảnh và cảm xúc.
STORY BIBLE là luật ưu tiên cao hơn suy đoán của bạn.

BẮT BUỘC:
- Giữ đúng tên nhân vật, giới tính và quan hệ.
- TUYỆT ĐỐI tuân thủ bảng xưng hô trong STORY BIBLE.
- Không tự ý đổi "ta" thành "tôi", "nàng" thành "cô ấy", "chàng" thành "anh ấy" nếu luật xưng hô không cho phép.
- Không hiện đại hóa văn phong cổ trang.
- Không thêm, bớt hoặc giải thích nội dung.
- Giữ đủ và đúng thứ tự các đoạn [P1], [P2]...
- Không đưa nhãn [P] vào bản dịch.

Trả JSON duy nhất:
{"paragraphs":["..."],"memory_updates":{"characters":[],"glossary":[],"world_rules":[],"style_rules":[]}}"""
    user=f"""CHƯƠNG: {title}

PHONG CÁCH:
{style}

BỘ NHỚ LIÊN QUAN:
{json.dumps(local,ensure_ascii=False,indent=2)}

ĐOẠN GỐC:
{source}
"""
    model=MODEL_FLASH if use_flash else MODEL_PRO
    d=call_json(client,model,system,user,thinking=False,timeout=180)
    paras=d.get("paragraphs",[])
    if len(paras)!=len(part):
        # Retry bằng Pro nếu Flash làm lệch cấu trúc.
        d=call_json(client,MODEL_PRO,system,user+f"\nPHẢI TRẢ ĐÚNG {len(part)} ĐOẠN.",thinking=False,timeout=240)
        paras=d.get("paragraphs",[])
    if len(paras)!=len(part):
        raise ValueError(f"Số đoạn đầu ra không khớp: {len(paras)}/{len(part)}")
    return paras,d.get("memory_updates",{})

def quality_check(client, title, source_parts, translated_parts, mem, style):
    source="\n".join(f"[P{i+1}] {p}" for i,p in enumerate(source_parts))
    target="\n".join(f"[P{i+1}] {p}" for i,p in enumerate(translated_parts))
    local=relevant_memory(mem,source)
    system="""Bạn là biên tập viên kiểm định bản dịch tiểu thuyết.
Chỉ tìm lỗi thực sự: sai tên, sai giới tính, sai quan hệ, sai xưng hô theo STORY BIBLE,
sai thuật ngữ, thiếu/đổi ý quan trọng hoặc văn phong lệch bối cảnh.
Không sửa chỉ vì sở thích.
Trả JSON:
{"pass":true/false,"issues":[{"type":"","original":"","current":"","fix":""}]}"""
    user=f"""CHƯƠNG: {title}
STYLE:
{style}
STORY BIBLE LIÊN QUAN:
{json.dumps(local,ensure_ascii=False,indent=2)}

NGUYÊN TÁC:
{source}

BẢN DỊCH:
{target}
"""
    return call_json(client,MODEL_PRO,system,user,thinking=False,timeout=240)

def repair(client,title,source_parts,translated_parts,check,mem,style):
    system="""Bạn là biên tập viên sửa bản dịch tiểu thuyết.
Chỉ sửa các lỗi được kiểm định; giữ nguyên mọi phần đúng.
Đặc biệt khóa xưng hô và tên theo STORY BIBLE.
Trả JSON {"paragraphs":["..."]} với đúng số đoạn đầu vào."""
    user=f"""CHƯƠNG: {title}
STYLE:
{style}
STORY BIBLE:
{json.dumps(relevant_memory(mem,' '.join(source_parts)),ensure_ascii=False,indent=2)}
LỖI ĐƯỢC XÁC ĐỊNH:
{json.dumps(check,ensure_ascii=False,indent=2)}
BẢN DỊCH:
{json.dumps(translated_parts,ensure_ascii=False)}
"""
    d=call_json(client,MODEL_PRO,system,user,thinking=False,timeout=240)
    if len(d.get("paragraphs",[]))!=len(translated_parts):
        raise ValueError("Bản sửa không giữ đúng số đoạn.")
    return d["paragraphs"]

def export_docx(book):
    doc=Document()
    for ch in book:
        doc.add_heading(ch["title"],level=1)
        for p in ch["paragraphs"]: doc.add_paragraph(p)
        if ch!=book[-1]: doc.add_page_break()
    b=io.BytesIO(); doc.save(b); return b.getvalue()

def export_txt(book):
    lines=[]
    for ch in book:
        lines += [ch["title"],""]+ch["paragraphs"]+["",""]
    return "\n".join(lines).encode("utf-8-sig")

st.sidebar.title("⚙️ Cài đặt")
api_key=st.sidebar.text_input("DeepSeek API Key",type="password")
mode=st.sidebar.radio("Chế độ",["⭐ Cân bằng","🧠 Chất lượng cao"],index=0)
style=st.sidebar.text_area("Phong cách dịch",value=DEFAULT_STYLE,height=200)
limit=st.sidebar.slider("Độ dài chunk",6000,12000,9000,500)
workers=st.sidebar.slider("Số lượt dịch song song",1,6,3)
check_all=st.sidebar.checkbox("Kiểm tra V4 Pro toàn bộ chunk",value=True)
st.sidebar.caption("Cân bằng: V4 Pro xây bộ nhớ + kiểm tra, V4 Flash dịch song song. Chất lượng cao: V4 Pro dịch và kiểm tra.")

st.title("📖 Dịch Truyện AI V6 — Cân bằng")
st.caption("Nhanh hơn nhưng vẫn giữ STORY BIBLE, xưng hô, tên riêng và kiểm tra bằng V4 Pro.")

file=st.file_uploader("📄 Tải file truyện tiếng Trung",type=["docx","txt"])

if file:
    raw=file.getvalue()
    paragraphs=read_docx(raw) if file.name.lower().endswith(".docx") else read_txt(raw)
    chapters=split_chapters(paragraphs)
    total_chunks=sum(len(split_smart(c["paragraphs"],limit)) for c in chapters)
    c1,c2,c3=st.columns(3)
    c1.metric("Số chương",len(chapters))
    c2.metric("Số đoạn",len(paragraphs))
    c3.metric("Số chunk",total_chunks)

    if st.button("🚀 Bắt đầu dịch",type="primary",disabled=not api_key):
        client=OpenAI(api_key=api_key,base_url="https://api.deepseek.com",timeout=240)

        left,right=st.columns([3,1])
        with right:
            st.markdown("### 📊 TIẾN ĐỘ")
            stage_box=st.empty()
            progress=st.progress(0)
            count_box=st.empty()
            current_box=st.empty()
            log_box=st.empty()

        def update(stage,done,total,current="",log=""):
            stage_box.markdown(f"**{stage}**")
            progress.progress(min(done/max(total,1),1.0))
            count_box.write(f"**{done}/{total}** lượt hoàn thành")
            if current: current_box.info(current)
            if log: log_box.write(log)

        update("1/4 — Xây STORY BIBLE",0,1,"Đang phân tích nhân vật, quan hệ, xưng hô...")
        mem=build_bible(client,chapters,style)
        st.session_state["memory"]=mem

        # Chuẩn bị toàn bộ chunk.
        tasks=[]
        for ch in chapters:
            parts=split_smart(ch["paragraphs"],limit)
            for idx,part in enumerate(parts):
                tasks.append((ch,idx,parts,part))
        total=len(tasks)
        results=[None]*total
        done=0
        update("2/4 — DỊCH",0,total)

        def run_task(item):
            ch,idx,parts,part=item
            tr,upd=translate_one(client,ch["title"],part,mem,style,use_flash=(mode=="⭐ Cân bằng"))
            return idx,tr,upd

        # Song song các chunk không phụ thuộc nhau. Memory updates được gom sau khi dịch.
        # Điều này tăng tốc; STORY BIBLE gốc vẫn được gửi cho từng request.
        max_workers=workers if mode=="⭐ Cân bằng" else min(2,workers)
        with concurrent.futures.ThreadPoolExecutor(max_workers=max_workers) as ex:
            futures={ex.submit(run_task,item):i for i,item in enumerate(tasks)}
            for fut in concurrent.futures.as_completed(futures):
                pos=futures[fut]
                ch,idx,parts,part=tasks[pos]
                try:
                    tr,upd=fut.result()
                    results[pos]=(tr,upd)
                    done+=1
                    update("2/4 — DỊCH",done,total,f"🔄 {ch['title']} — chunk {idx+1}/{len(parts)}",
                           f"Đã hoàn thành {ch['title']} chunk {idx+1}")
                except Exception as e:
                    results[pos]=(["[LỖI API — CẦN DỊCH LẠI ĐOẠN NÀY]"]*len(part),{})
                    done+=1
                    update("2/4 — DỊCH",done,total,f"⚠️ {ch['title']} — chunk {idx+1}/{len(parts)}",str(e))

        # Gom memory updates theo thứ tự file để ổn định.
        for tr_upd in results:
            if tr_upd: mem=merge_memory(mem,tr_upd[1])
        st.session_state["memory"]=mem

        # Ghép theo chương.
        book=[]
        pos=0
        for ch in chapters:
            parts=split_smart(ch["paragraphs"],limit)
            paras=[]
            for _ in parts:
                paras.extend(results[pos][0]); pos+=1
            book.append({"title":f"Chương {ch['number']}","paragraphs":paras})

        if check_all:
            update("3/4 — KIỂM TRA V4 PRO",0,total,"Đang kiểm tra tên, xưng hô, thuật ngữ...")
            checked=0
            pos=0
            for bi,ch in enumerate(chapters):
                parts=split_smart(ch["paragraphs"],limit)
                rebuilt=[]
                for part in parts:
                    translated=results[pos][0]
                    try:
                        ck=quality_check(client,ch["title"],part,translated,mem,style)
                        if not ck.get("pass",True):
                            translated=repair(client,ch["title"],part,translated,ck,mem,style)
                    except Exception:
                        pass
                    rebuilt.extend(translated)
                    checked+=1; pos+=1
                    update("3/4 — KIỂM TRA V4 PRO",checked,total,f"🔍 {ch['title']} — chunk {checked}",
                           f"Kiểm tra {checked}/{total}")
                book[bi]["paragraphs"]=rebuilt

        update("4/4 — HOÀN TẤT",total,total,"🎉 Đã hoàn thành và ghép file Word.")
        st.session_state["book"]=book

if "memory" in st.session_state:
    with st.expander("🧠 STORY BIBLE — Bộ nhớ nhân vật/xưng hô"):
        st.json(st.session_state["memory"])

if "book" in st.session_state:
    st.subheader("📥 Tải bản dịch")
    st.download_button("⬇️ TẢI FILE WORD ĐÃ DỊCH",export_docx(st.session_state["book"]),
                       "ban-dich-truyen-v6.docx",
                       "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                       type="primary",use_container_width=True)
    st.download_button("⬇️ Tải TXT",export_txt(st.session_state["book"]),
                       "ban-dich-truyen-v6.txt","text/plain",use_container_width=True)
